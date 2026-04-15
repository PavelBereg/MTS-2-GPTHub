"""
DOCX Generator — эндпоинт для генерации документов Word.

Маршруты:
  POST /v1/documents/generate             — создаёт DOCX по теме
  POST /v1/documents/generate-from-plan   — создаёт DOCX из согласованного плана
  GET  /v1/documents/download/{file_id}   — скачивает готовый файл
  GET  /v1/documents/list/{user_id}       — список документов пользователя

Интеграция в чат:
  document-agent сначала собирает бриф, предлагает структуру и только после
  подтверждения пользователя передаёт план сюда для генерации текста и DOCX.
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
import time
import uuid
from pathlib import Path
from typing import Any

import httpx
from fastapi import APIRouter
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field

from app.core.config import get_settings

logger = logging.getLogger("mts.docx")
settings = get_settings()
router = APIRouter()

DOCX_DIR = Path("/tmp/mts_documents")
DOCX_DIR.mkdir(parents=True, exist_ok=True)

_file_registry: dict[str, dict[str, Any]] = {}


DOCUMENT_STYLE_SPECS: dict[str, dict[str, Any]] = {
    "official_business": {
        "label": "официальный деловой документ",
        "font": "Times New Roman",
        "font_size": 14,
        "line_spacing": 1.5,
        "alignment": "justify",
        "first_line_indent_cm": 1.25,
        "margins_cm": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
        "tone": "строго формальный, канцелярский, без эмоций",
        "structure": "заголовок по центру, основной текст, заключение, подпись/дата при необходимости",
        "numbering": "1., 1.1., 1.1.1.; маркированные списки только если без них хуже",
    },
    "academic": {
        "label": "академическая работа",
        "font": "Times New Roman",
        "font_size": 14,
        "line_spacing": 1.5,
        "alignment": "justify",
        "first_line_indent_cm": 1.25,
        "margins_cm": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
        "tone": "научный, доказательный, без разговорной лексики",
        "structure": "титульный лист, содержание, введение, главы, заключение, список литературы",
        "numbering": "1 ГЛАВА, 1.1 Подраздел, 1.1.1 Подпункт; ссылки вида [1], [2, с. 45]",
    },
    "business_report": {
        "label": "бизнес-документ",
        "font": "Calibri",
        "font_size": 11,
        "line_spacing": 1.15,
        "alignment": "left",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.2, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "ясный, короткий, управленческий",
        "structure": "заголовок, executive summary, основная часть, выводы, рекомендации",
        "numbering": "маркированные списки, короткие блоки, таблицы для сравнений",
    },
    "business_letter": {
        "label": "деловое письмо",
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.15,
        "alignment": "justify",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.5, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "вежливый, краткий, адресный",
        "structure": "кому, от кого, дата, обращение, текст письма, подпись, контакты",
        "numbering": "сложная нумерация не нужна",
    },
    "legal": {
        "label": "юридический документ",
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "alignment": "justify",
        "first_line_indent_cm": 1.25,
        "margins_cm": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
        "tone": "максимально точный, без разговорных слов, с однозначными формулировками",
        "structure": "предмет, термины, права и обязанности, порядок действия, ответственность, подписи",
        "numbering": "1., 1.1., 1.1.1.; каждый пункт должен быть юридически однозначным",
    },
    "marketing": {
        "label": "маркетинговый текст / презентационный документ",
        "font": "Arial",
        "font_size": 12,
        "line_spacing": 1.15,
        "alignment": "left",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "живой, продающий, простой, но без пустых обещаний",
        "structure": "заголовок, сильное резюме, выгоды, доказательства, призыв к действию",
        "numbering": "короткие абзацы, жирные акценты, списки и таблицы выгод",
    },
    "informal_notes": {
        "label": "неформальный документ / заметки",
        "font": "Calibri",
        "font_size": 11,
        "line_spacing": 1.0,
        "alignment": "left",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "свободный, понятный, допускаются сокращения и лёгкая разговорность",
        "structure": "короткие блоки, списки, заметки, выводы",
        "numbering": "любой удобный формат списков",
    },
    "technical_docs": {
        "label": "техническая документация",
        "font": "Calibri",
        "font_size": 11,
        "code_font": "Consolas",
        "line_spacing": 1.15,
        "alignment": "left",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "точный, структурный, без маркетинговой воды",
        "structure": "описание, требования, архитектура/логика, инструкции, примеры, ограничения",
        "numbering": "много списков, шаги, таблицы параметров, код отдельными блоками",
    },
    "instruction": {
        "label": "инструкция / руководство",
        "font": "Calibri",
        "font_size": 11,
        "line_spacing": 1.15,
        "alignment": "left",
        "first_line_indent_cm": 0.0,
        "margins_cm": {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0},
        "tone": "пошаговый, простой, в повелительном наклонении",
        "structure": "цель, перед началом, шаги, проверка результата, частые ошибки",
        "numbering": "1. Откройте; 2. Нажмите; 3. Проверьте",
    },
}


class DocumentRequest(BaseModel):
    topic: str = Field(..., min_length=3, max_length=1000)
    document_type: str = Field(default="business_report")
    audience: str = Field(default="")
    goal: str = Field(default="")
    volume_pages: int = Field(default=3, ge=1, le=25)
    language: str = Field(default="ru")
    source_material: str = Field(default="")
    user_id: str = Field(default="anonymous")


class DocumentPlanSectionRequest(BaseModel):
    heading: str
    level: int = Field(default=1, ge=1, le=3)
    purpose: str = ""
    key_points: list[str] = Field(default_factory=list)
    facts_to_highlight: list[str] = Field(default_factory=list)
    format_hint: str = ""
    expected_elements: list[str] = Field(default_factory=list)


class DocumentFromPlanRequest(BaseModel):
    brief: dict[str, Any] = Field(default_factory=dict)
    plan: list[DocumentPlanSectionRequest] = Field(default_factory=list)
    document_type: str = Field(default="business_report")
    language: str = Field(default="ru")
    user_id: str = Field(default="anonymous")


class DocumentTable(BaseModel):
    headers: list[str] = Field(default_factory=list)
    rows: list[list[str]] = Field(default_factory=list)


class DocumentSection(BaseModel):
    heading: str
    level: int = Field(default=1, ge=1, le=3)
    paragraphs: list[str] = Field(default_factory=list)
    bullets: list[str] = Field(default_factory=list)
    numbered_items: list[str] = Field(default_factory=list)
    facts: list[str] = Field(default_factory=list)
    table: DocumentTable | None = None
    code_blocks: list[str] = Field(default_factory=list)
    callout: str = ""


class DocumentData(BaseModel):
    title: str
    subtitle: str = ""
    document_type: str = "business_report"
    add_title_page: bool = False
    summary: list[str] = Field(default_factory=list)
    metadata_lines: list[str] = Field(default_factory=list)
    sections: list[DocumentSection]
    conclusion: list[str] = Field(default_factory=list)
    references: list[str] = Field(default_factory=list)
    signature_block: list[str] = Field(default_factory=list)


def _shorten_text(text: Any, limit: int) -> str:
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(value) <= limit:
        return value
    cut = value[: max(0, limit - 1)].rstrip()
    if " " in cut:
        cut = cut.rsplit(" ", 1)[0]
    return cut.rstrip(".,;:") + "..."


def _strip_think_and_fences(text: str) -> str:
    text = re.sub(r"<think>[\s\S]*?</think>", "", text or "").strip()
    text = re.sub(r"```(?:json)?\s*", "", text).strip("` \n")
    return text


def _extract_json_object(text: str) -> dict[str, Any]:
    raw = _strip_think_and_fences(text)
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"LLM не вернул JSON. Ответ: {raw[:300]}")
    try:
        obj = json.loads(raw[start : end + 1])
    except json.JSONDecodeError as e:
        raise ValueError(f"Невалидный JSON от LLM: {e}. Raw: {raw[:300]}") from e
    if not isinstance(obj, dict):
        raise ValueError("LLM вернул JSON не-объект.")
    return obj


def normalize_document_type(value: Any) -> str:
    t = str(value or "").lower().strip()
    if not t:
        return "business_report"
    if any(k in t for k in ("гост", "официаль", "служеб", "деловой документ", "приказ", "заявлен")):
        return "official_business"
    if any(k in t for k in ("академ", "реферат", "курсов", "диплом", "науч", "доклад")):
        return "academic"
    if any(k in t for k in ("бизнес", "отч", "proposal", "предлож", "executive", "аналитическ")):
        return "business_report"
    if any(k in t for k in ("письм", "кому", "с уважением", "email", "e-mail")):
        return "business_letter"
    if any(k in t for k in ("юрид", "договор", "соглаш", "оферт", "стороны", "правов")):
        return "legal"
    if any(k in t for k in ("маркет", "прода", "лендинг", "коммерческ", "презентацион")):
        return "marketing"
    if any(k in t for k in ("неформ", "замет", "черновик", "свободн")):
        return "informal_notes"
    if any(k in t for k in ("тех", "api", "архитект", "документац", "тз", "техническ")):
        return "technical_docs"
    if any(k in t for k in ("инструк", "руковод", "гайд", "manual", "пошаг")):
        return "instruction"
    if t in DOCUMENT_STYLE_SPECS:
        return t
    return "business_report"


def _style_spec_for(document_type: str) -> dict[str, Any]:
    normalized = normalize_document_type(document_type)
    return DOCUMENT_STYLE_SPECS.get(normalized, DOCUMENT_STYLE_SPECS["business_report"])


_CONTENT_PROMPT = """\
Ты — senior document writer и редактор Word-документов. Сгенерируй полноценный
документ по согласованному брифу и плану.

Бриф:
{brief_json}

Согласованный план:
{plan_json}

Стандарт оформления выбранного стиля:
{style_json}

Верни ТОЛЬКО валидный JSON без markdown:
{{
  "title": "короткий заголовок документа",
  "subtitle": "подзаголовок или пустая строка",
  "document_type": "{document_type}",
  "add_title_page": true|false,
  "metadata_lines": ["Кому: ...", "От кого: ...", "Дата: ..."],
  "summary": ["краткий вывод или executive summary", "..."],
  "sections": [
    {{
      "heading": "Название раздела",
      "level": 1,
      "paragraphs": ["полноценный абзац", "ещё абзац"],
      "bullets": ["маркированный пункт"],
      "numbered_items": ["нумерованный пункт"],
      "facts": ["факт из материала или помеченный вывод"],
      "table": {{"headers": ["Параметр", "Описание"], "rows": [["...", "..."]]}},
      "code_blocks": ["пример кода, если нужен"],
      "callout": "важное замечание или пустая строка"
    }}
  ],
  "conclusion": ["итоговый вывод", "следующее действие"],
  "references": ["[1] источник или пометка, если пользователь дал источник"],
  "signature_block": ["С уважением,", "Имя Фамилия"]
}}

Жёсткие правила качества:
- Следуй выбранному типу документа, его структуре, тону и стандартам оформления.
- Не генерируй презентационные слайды. Это именно Word-документ с развёрнутыми абзацами.
- Если пользователь дал source_material, опирайся прежде всего на него и не теряй факты.
- Факты отделяй от выводов: не выдумывай числа, даты, имена, законы и источники.
- Если данных не хватает, формулируй аккуратно: "требует проверки", "по предоставленным данным".
- Для официального и юридического стиля используй точные формулировки и нумерацию.
- Для академического стиля добавь введение, основную часть, заключение и список литературы;
  ссылки оформляй в квадратных скобках, но не придумывай реальные источники.
- Для бизнес-документа обязательно дай summary, выводы и рекомендации.
- Для технической документации используй требования, инструкции, таблицы параметров и code_blocks только по делу.
- Для инструкции пиши шаги в императиве: "Откройте", "Нажмите", "Проверьте".
- В каждом содержательном разделе должны быть конкретика, причинно-следственные связи и вывод.
- Не используй пустые фразы вроде "данный документ является важным" без объяснения почему.
- Объём должен соответствовать brief.volume_pages: не короче рабочего черновика, но без воды.
"""


async def _generate_document_data_from_plan(
    request: DocumentFromPlanRequest,
) -> DocumentData:
    document_type = normalize_document_type(
        request.document_type or request.brief.get("document_type")
    )
    brief = dict(request.brief or {})
    brief["document_type"] = document_type
    spec = _style_spec_for(document_type)

    prompt = _CONTENT_PROMPT.format(
        brief_json=json.dumps(brief, ensure_ascii=False),
        plan_json=json.dumps([s.model_dump() for s in request.plan], ensure_ascii=False),
        style_json=json.dumps(spec, ensure_ascii=False),
        document_type=document_type,
    )

    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(30.0, read=120.0)) as client:
            response = await client.post(
                f"{settings.MWS_BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings.MWS_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": settings.LLM_MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.35,
                    "max_tokens": 4096,
                    "stream": False,
                },
            )
            response.raise_for_status()
            raw = response.json()["choices"][0]["message"]["content"]
    except httpx.TimeoutException as e:
        raise ValueError(f"LLM timeout при генерации документа: {e}") from e
    except Exception as e:
        raise ValueError(f"LLM API error: {e}") from e

    obj = _extract_json_object(str(raw))
    obj["document_type"] = document_type
    return _normalize_document_data(obj, brief, request.plan)


def _normalize_document_data(
    raw: dict[str, Any],
    brief: dict[str, Any],
    plan: list[DocumentPlanSectionRequest],
) -> DocumentData:
    sections: list[DocumentSection] = []
    for idx, item in enumerate(raw.get("sections", []), 1):
        if not isinstance(item, dict):
            continue
        table = None
        table_raw = item.get("table")
        if isinstance(table_raw, dict):
            headers = [str(h).strip() for h in table_raw.get("headers", []) if str(h).strip()]
            rows = []
            for row in table_raw.get("rows", []):
                if isinstance(row, list):
                    cleaned = [str(cell).strip() for cell in row]
                    if any(cleaned):
                        rows.append(cleaned[: max(1, len(headers) or len(cleaned))])
            if headers and rows:
                table = DocumentTable(headers=headers[:6], rows=rows[:12])

        sections.append(
            DocumentSection(
                heading=_shorten_text(item.get("heading") or f"Раздел {idx}", 120),
                level=max(1, min(int(item.get("level") or 1), 3)),
                paragraphs=[
                    str(p).strip()
                    for p in item.get("paragraphs", [])
                    if str(p).strip()
                ][:8]
                if isinstance(item.get("paragraphs", []), list)
                else [],
                bullets=[
                    _shorten_text(b, 240)
                    for b in item.get("bullets", [])
                    if str(b).strip()
                ][:10]
                if isinstance(item.get("bullets", []), list)
                else [],
                numbered_items=[
                    _shorten_text(n, 260)
                    for n in item.get("numbered_items", [])
                    if str(n).strip()
                ][:12]
                if isinstance(item.get("numbered_items", []), list)
                else [],
                facts=[
                    _shorten_text(f, 260)
                    for f in item.get("facts", [])
                    if str(f).strip()
                ][:8]
                if isinstance(item.get("facts", []), list)
                else [],
                table=table,
                code_blocks=[
                    str(c).strip()
                    for c in item.get("code_blocks", [])
                    if str(c).strip()
                ][:4]
                if isinstance(item.get("code_blocks", []), list)
                else [],
                callout=_shorten_text(item.get("callout") or "", 320),
            )
        )

    if not sections:
        for idx, item in enumerate(plan, 1):
            points = item.key_points or item.facts_to_highlight or ["Раскрыть ключевой аспект темы."]
            sections.append(
                DocumentSection(
                    heading=item.heading or f"Раздел {idx}",
                    level=item.level,
                    paragraphs=[
                        " ".join(points[:3])
                        + " Этот раздел требует редакторской проверки и уточнения фактов."
                    ],
                    facts=item.facts_to_highlight[:5],
                )
            )

    title = str(raw.get("title") or brief.get("title") or brief.get("topic") or "Документ").strip()
    doc_type = normalize_document_type(raw.get("document_type") or brief.get("document_type"))
    return DocumentData(
        title=_shorten_text(title, 140),
        subtitle=_shorten_text(raw.get("subtitle") or "", 180),
        document_type=doc_type,
        add_title_page=bool(raw.get("add_title_page", doc_type == "academic")),
        metadata_lines=[
            _shorten_text(line, 180)
            for line in raw.get("metadata_lines", [])
            if str(line).strip()
        ][:10]
        if isinstance(raw.get("metadata_lines", []), list)
        else [],
        summary=[
            _shorten_text(item, 320)
            for item in raw.get("summary", [])
            if str(item).strip()
        ][:8]
        if isinstance(raw.get("summary", []), list)
        else [],
        sections=sections,
        conclusion=[
            _shorten_text(item, 320)
            for item in raw.get("conclusion", [])
            if str(item).strip()
        ][:8]
        if isinstance(raw.get("conclusion", []), list)
        else [],
        references=[
            _shorten_text(item, 260)
            for item in raw.get("references", [])
            if str(item).strip()
        ][:20]
        if isinstance(raw.get("references", []), list)
        else [],
        signature_block=[
            _shorten_text(item, 160)
            for item in raw.get("signature_block", [])
            if str(item).strip()
        ][:8]
        if isinstance(raw.get("signature_block", []), list)
        else [],
    )


def _safe_filename(title: str, file_id: str) -> str:
    value = re.sub(r"[^\w\s-]", "", title, flags=re.UNICODE)[:55].strip()
    value = re.sub(r"\s+", "_", value)
    return f"{value or 'document'}_{file_id}.docx"


def _build_docx(data: DocumentData, document_type: str, file_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as e:
        raise RuntimeError(
            "python-docx не установлен. Добавьте 'python-docx' в requirements.txt."
        ) from e

    doc_type = normalize_document_type(document_type or data.document_type)
    spec = _style_spec_for(doc_type)
    document = Document()

    section = document.sections[0]
    margins = spec["margins_cm"]
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])

    normal = document.styles["Normal"]
    normal.font.name = spec["font"]
    normal.font.size = Pt(spec["font_size"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), spec["font"])
    normal.paragraph_format.line_spacing = spec["line_spacing"]
    normal.paragraph_format.first_line_indent = Cm(spec["first_line_indent_cm"])

    for style_name, size, bold in (
        ("Title", spec["font_size"] + 4, True),
        ("Heading 1", spec["font_size"] + 2, True),
        ("Heading 2", spec["font_size"] + 1, True),
        ("Heading 3", spec["font_size"], True),
    ):
        style = document.styles[style_name]
        style.font.name = spec["font"]
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), spec["font"])

    def _alignment(value: str):
        if value == "justify":
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        if value == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if value == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.LEFT

    body_alignment = _alignment(spec["alignment"])

    def _set_paragraph_format(paragraph, *, first_line: bool = True, align=None):
        paragraph.alignment = body_alignment if align is None else align
        paragraph.paragraph_format.line_spacing = spec["line_spacing"]
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.first_line_indent = (
            Cm(spec["first_line_indent_cm"]) if first_line else Cm(0)
        )

    def _add_run_paragraph(text: str, *, bold: bool = False, style: str | None = None):
        paragraph = document.add_paragraph(style=style)
        run = paragraph.add_run(str(text).strip())
        run.bold = bold
        run.font.name = spec["font"]
        run.font.size = Pt(spec["font_size"])
        _set_paragraph_format(paragraph, first_line=not style)
        return paragraph

    def _shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _add_table(table_data: DocumentTable) -> None:
        headers = table_data.headers[:6]
        rows = table_data.rows[:12]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            _shade_cell(cell, "F2F2F2")
        for row in rows:
            cells = table.add_row().cells
            for idx, cell in enumerate(cells):
                cell.text = str(row[idx] if idx < len(row) else "")
        document.add_paragraph()

    if data.add_title_page and doc_type == "academic":
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(180)
        run = title.add_run(data.title.upper())
        run.bold = True
        run.font.name = spec["font"]
        run.font.size = Pt(spec["font_size"] + 2)
        if data.subtitle:
            sub = document.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.add_run(data.subtitle)
        document.add_page_break()

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title.add_run(data.title)

    if data.subtitle:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.first_line_indent = Cm(0)
        run = subtitle.add_run(data.subtitle)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for line in data.metadata_lines:
        _add_run_paragraph(line, style=None)

    if data.summary:
        heading = "Executive Summary" if doc_type == "business_report" else "Краткое резюме"
        document.add_heading(heading, level=1)
        for item in data.summary:
            _add_run_paragraph(item)

    for section_data in data.sections:
        document.add_heading(section_data.heading, level=section_data.level)
        for paragraph in section_data.paragraphs:
            if paragraph.strip():
                _add_run_paragraph(paragraph)
        if section_data.facts:
            facts_title = "Факты и основания" if doc_type in {"official_business", "legal"} else "Факты"
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(f"{facts_title}:").bold = True
            for fact in section_data.facts:
                item = document.add_paragraph(style="List Bullet")
                item.add_run(fact)
                _set_paragraph_format(item, first_line=False, align=body_alignment)
        for bullet in section_data.bullets:
            item = document.add_paragraph(style="List Bullet")
            item.add_run(bullet)
            _set_paragraph_format(item, first_line=False, align=body_alignment)
        for numbered in section_data.numbered_items:
            item = document.add_paragraph(style="List Number")
            item.add_run(numbered)
            _set_paragraph_format(item, first_line=False, align=body_alignment)
        if section_data.table:
            _add_table(section_data.table)
        for code in section_data.code_blocks:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.5)
            run = paragraph.add_run(code)
            run.font.name = spec.get("code_font", "Consolas")
            run.font.size = Pt(max(9, spec["font_size"] - 1))
        if section_data.callout:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run("Важно: " + section_data.callout)
            run.bold = True

    if data.conclusion:
        document.add_heading("Заключение", level=1)
        for item in data.conclusion:
            _add_run_paragraph(item)

    if data.references:
        document.add_heading("Список литературы" if doc_type == "academic" else "Источники и примечания", level=1)
        for ref in data.references:
            item = document.add_paragraph(style="List Number")
            item.add_run(ref)
            _set_paragraph_format(item, first_line=False, align=body_alignment)

    if data.signature_block:
        document.add_paragraph()
        for line in data.signature_block:
            _add_run_paragraph(line)

    document.save(str(file_path))
    logger.info("DOCX: saved to %s (%d bytes)", file_path, file_path.stat().st_size)


def _preview_markdown(data: DocumentData) -> str:
    lines = [
        f"## 📄 {data.title}",
        f"*{data.subtitle}*" if data.subtitle else "",
        "",
        "### Структура:",
    ]
    for idx, section in enumerate(data.sections, 1):
        lines.append(f"**{idx}. {section.heading}**")
        if section.paragraphs:
            lines.append(f"  - {_shorten_text(section.paragraphs[0], 120)}")
        elif section.bullets:
            lines.append(f"  - {_shorten_text(section.bullets[0], 120)}")
    if data.conclusion:
        lines += ["", "### Вывод:", f"> {_shorten_text(data.conclusion[0], 220)}"]
    return "\n".join(filter(None, lines))


async def _render_and_register_docx(
    data: DocumentData,
    *,
    topic: str,
    document_type: str,
    user_id: str,
) -> JSONResponse:
    file_id = uuid.uuid4().hex[:16]
    filename = _safe_filename(data.title, file_id)
    file_path = DOCX_DIR / filename

    try:
        await asyncio.to_thread(_build_docx, data, document_type, file_path)
    except RuntimeError as e:
        logger.error("DOCX build failed: %s", e)
        return JSONResponse(status_code=500, content={"error": str(e)})
    except Exception:
        logger.exception("DOCX build unexpected error")
        return JSONResponse(
            status_code=500,
            content={"error": "Внутренняя ошибка создания DOCX-файла."},
        )

    _file_registry[file_id] = {
        "path": str(file_path),
        "filename": filename,
        "user_id": user_id,
        "topic": topic,
        "document_type": normalize_document_type(document_type),
        "created_at": time.time(),
        "section_count": len(data.sections),
    }

    download_url = f"/v1/documents/download/{file_id}"
    logger.info(
        "DOCX ready: file_id=%s sections=%d size=%dKB",
        file_id,
        len(data.sections),
        file_path.stat().st_size // 1024,
    )

    return JSONResponse(
        content={
            "file_id": file_id,
            "download_url": download_url,
            "filename": filename,
            "section_count": len(data.sections),
            "topic": topic,
            "document_type": normalize_document_type(document_type),
            "preview_markdown": _preview_markdown(data),
        }
    )


def _request_to_plan(request: DocumentRequest) -> DocumentFromPlanRequest:
    document_type = normalize_document_type(request.document_type)
    brief = {
        "topic": request.topic,
        "document_type": document_type,
        "audience": request.audience,
        "goal": request.goal,
        "volume_pages": request.volume_pages,
        "language": request.language,
        "source_material": request.source_material,
    }
    default_sections = {
        "official_business": ["Основание", "Основные положения", "Заключительные положения"],
        "academic": ["Введение", "Основная часть", "Заключение", "Список литературы"],
        "business_report": ["Executive Summary", "Контекст", "Анализ", "Рекомендации"],
        "business_letter": ["Обращение", "Суть письма", "Заключение"],
        "legal": ["Предмет", "Права и обязанности сторон", "Порядок действия", "Ответственность"],
        "marketing": ["Ключевое предложение", "Выгоды", "Доказательства", "Призыв к действию"],
        "informal_notes": ["Контекст", "Заметки", "Выводы"],
        "technical_docs": ["Описание", "Требования", "Инструкция", "Примеры"],
        "instruction": ["Перед началом", "Шаги", "Проверка результата", "Частые ошибки"],
    }
    plan = [
        DocumentPlanSectionRequest(
            heading=heading,
            level=1,
            purpose=f"Раскрыть блок «{heading}» по теме документа.",
            key_points=[request.topic],
            facts_to_highlight=[],
            format_hint=_style_spec_for(document_type)["structure"],
            expected_elements=["paragraphs", "bullets"],
        )
        for heading in default_sections.get(document_type, default_sections["business_report"])
    ]
    return DocumentFromPlanRequest(
        brief=brief,
        plan=plan,
        document_type=document_type,
        language=request.language,
        user_id=request.user_id,
    )


@router.post("/documents/generate")
async def generate_document(request: DocumentRequest):
    """Генерирует DOCX-документ по теме без интерактивного брифинга."""
    plan_request = _request_to_plan(request)
    try:
        data = await _generate_document_data_from_plan(plan_request)
    except ValueError as e:
        logger.error("DOCX content generation failed: %s", e)
        return JSONResponse(
            status_code=502,
            content={"error": f"Ошибка генерации документа: {str(e)[:300]}"},
        )

    return await _render_and_register_docx(
        data,
        topic=request.topic,
        document_type=plan_request.document_type,
        user_id=request.user_id,
    )


@router.post("/documents/generate-from-plan")
async def generate_document_from_plan(request: DocumentFromPlanRequest):
    """Генерирует DOCX из уже согласованного брифа и плана документа."""
    logger.info(
        "documents/generate-from-plan: sections=%d type=%s user=%s",
        len(request.plan),
        request.document_type,
        request.user_id,
    )
    if not request.plan:
        return JSONResponse(
            status_code=400,
            content={"error": "Нет согласованного плана документа для генерации DOCX."},
        )

    try:
        data = await _generate_document_data_from_plan(request)
    except ValueError as e:
        logger.error("DOCX content generation failed: %s", e)
        return JSONResponse(
            status_code=502,
            content={"error": f"Ошибка генерации документа: {str(e)[:300]}"},
        )

    topic = str(request.brief.get("topic") or data.title)
    return await _render_and_register_docx(
        data,
        topic=topic,
        document_type=request.document_type or data.document_type,
        user_id=request.user_id,
    )


@router.get("/documents/download/{file_id}")
async def download_document(file_id: str):
    """Скачивает готовый DOCX-файл по ID."""
    info = _file_registry.get(file_id)
    if not info:
        return JSONResponse(status_code=404, content={"error": "Файл не найден."})

    path = Path(info["path"])
    if not path.exists():
        return JSONResponse(
            status_code=410,
            content={"error": "Файл был удалён с сервера."},
        )

    return FileResponse(
        path=str(path),
        filename=info["filename"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/documents/list/{user_id}")
async def list_documents(user_id: str):
    """Возвращает список DOCX-документов пользователя."""
    user_files = [
        {
            "file_id": fid,
            "filename": info["filename"],
            "topic": info["topic"],
            "document_type": info["document_type"],
            "section_count": info["section_count"],
            "download_url": f"/v1/documents/download/{fid}",
            "created_at": info["created_at"],
        }
        for fid, info in _file_registry.items()
        if info.get("user_id") == user_id
    ]
    return JSONResponse(content={"documents": user_files, "count": len(user_files)})
